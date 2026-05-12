"""
ingestion.py — Document parsing, chunking, and embedding for Tome

Chunking strategy:
  .json  → each object = one chunk (guaranteed, no splitting)
  .md    → heading-aware: one chunk per ## section
  .docx  → heading-aware: one chunk per Word heading style
  .xlsx  → each row = one chunk
  .pdf   → heading-aware fallback
"""

import os
import re
import json
import pickle
import numpy as np

DATA_DIR          = os.path.join(os.path.dirname(__file__), "data")
FAISS_INDEX_PATH  = os.path.join(DATA_DIR, "faiss_index", "index.faiss")
EMBEDDINGS_PATH   = os.path.join(DATA_DIR, "faiss_index", "embeddings.pkl")

_model        = None
_faiss_index  = None
_embedding_store: list[np.ndarray] = []


def _get_model():
    global _model
    if _model is None:
        from sentence_transformers import SentenceTransformer
        _model = SentenceTransformer("all-MiniLM-L6-v2")
    return _model


def _get_faiss_index(dim: int = 384):
    global _faiss_index, _embedding_store
    import faiss
    os.makedirs(os.path.dirname(FAISS_INDEX_PATH), exist_ok=True)
    if os.path.exists(FAISS_INDEX_PATH):
        _faiss_index = faiss.read_index(FAISS_INDEX_PATH)
        if os.path.exists(EMBEDDINGS_PATH):
            with open(EMBEDDINGS_PATH, "rb") as f:
                _embedding_store = pickle.load(f)
    else:
        _faiss_index = faiss.IndexFlatL2(dim)
        _embedding_store = []
    return _faiss_index


def _save_index():
    import faiss
    if _faiss_index is not None:
        faiss.write_index(_faiss_index, FAISS_INDEX_PATH)
        with open(EMBEDDINGS_PATH, "wb") as f:
            pickle.dump(_embedding_store, f)


def get_index_size() -> int:
    return _get_faiss_index().ntotal


# ══════════════════════════════════════════════════════════════════════════════
# JSON PARSER — one object = one chunk
# ══════════════════════════════════════════════════════════════════════════════

def _chunks_from_json(file_bytes: bytes) -> list[str]:
    data = json.loads(file_bytes.decode("utf-8", errors="ignore"))
    if isinstance(data, dict):
        data = [data]

    chunks = []
    for item in data:
        if isinstance(item, str):
            if item.strip():
                chunks.append(item.strip())
            continue
        if isinstance(item, dict):
            q_val, a_val, other = "", "", []
            for k, v in item.items():
                kl  = k.lower()
                val = str(v).strip() if v is not None else ""
                if not val:
                    continue
                if any(x in kl for x in ["question", "title", "topic"]):
                    q_val = val
                elif any(x in kl for x in ["answer", "content", "body", "text"]):
                    a_val = val
                else:
                    other.append(f"{k}: {val}")
            parts = []
            if q_val:
                parts.append(f"Q: {q_val}")
            if a_val:
                parts.append(f"A: {a_val}")
            parts.extend(other)
            chunk = "\n".join(parts).strip()
            if chunk:
                chunks.append(chunk)
    return chunks


# ══════════════════════════════════════════════════════════════════════════════
# MARKDOWN PARSER — one chunk per ## heading block
# ══════════════════════════════════════════════════════════════════════════════

def _chunks_from_markdown(text: str) -> list[str]:
    chunks   = []
    lines    = text.split("\n")
    section  = ""
    buf_head = ""
    buf_body: list[str] = []

    def flush():
        body = "\n".join(buf_body).strip()
        head = buf_head.strip()
        ctx  = f"[{section}] " if section else ""
        if head and body:
            chunks.append(f"{ctx}{head}\n{body}")
        elif head:
            chunks.append(f"{ctx}{head}")
        elif body:
            chunks.append(f"{ctx}{body}")

    for line in lines:
        stripped = line.strip()
        if re.match(r"^-{3,}$", stripped):
            flush(); buf_head, buf_body = "", []; continue
        if re.match(r"^#\s+", line) and not re.match(r"^##", line):
            flush(); buf_head, buf_body = "", []
            section = line.lstrip("#").strip(); continue
        if re.match(r"^#{2,}\s+", line):
            flush()
            buf_head = line.lstrip("#").strip(); buf_body = []; continue
        if stripped:
            buf_body.append(stripped)

    flush()
    if not chunks:
        chunks = _word_window_chunks(text)
    return [c for c in chunks if len(c.strip()) > 15]


# ══════════════════════════════════════════════════════════════════════════════
# DOCX PARSER — one chunk per Word heading style
# ══════════════════════════════════════════════════════════════════════════════

def _chunks_from_docx(file_bytes: bytes) -> list[str]:
    import io
    from docx import Document
    doc      = Document(io.BytesIO(file_bytes))
    chunks   = []
    section  = ""
    buf_head = ""
    buf_body: list[str] = []

    def flush():
        body = " ".join(buf_body).strip()
        ctx  = f"[{section}] " if section else ""
        if buf_head and body:
            chunks.append(f"{ctx}{buf_head}\n{body}")
        elif buf_head:
            chunks.append(f"{ctx}{buf_head}")
        elif body:
            chunks.append(f"{ctx}{body}")

    for para in doc.paragraphs:
        text  = para.text.strip()
        style = para.style.name.lower() if para.style else ""
        if not text:
            continue
        if "heading 1" in style:
            flush(); buf_head, buf_body = "", []; section = text
        elif "heading 2" in style:
            flush(); buf_head = text; buf_body = []
        elif "heading 3" in style:
            buf_body.append(f"{text}:")
        else:
            buf_body.append(text)

    flush()
    if not chunks:
        plain  = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        chunks = _chunks_from_markdown(plain)
    return [c for c in chunks if len(c.strip()) > 15]


# ══════════════════════════════════════════════════════════════════════════════
# XLSX PARSER — one row = one chunk
# ══════════════════════════════════════════════════════════════════════════════

def _chunks_from_xlsx(file_bytes: bytes) -> list[str]:
    import io
    import pandas as pd
    xl     = pd.ExcelFile(io.BytesIO(file_bytes))
    chunks = []
    for sheet in xl.sheet_names:
        df   = xl.parse(sheet).fillna("")
        cols = [str(c).strip() for c in df.columns]
        q_col = next((c for c in cols if any(x in c.lower() for x in ["question", "title"])), None)
        a_col = next((c for c in cols if any(x in c.lower() for x in ["answer", "content", "body"])), None)
        for _, row in df.iterrows():
            values = {c: str(row[c]).strip() for c in cols if str(row[c]).strip()}
            if not values:
                continue
            if q_col and a_col and q_col in values and a_col in values:
                chunk  = f"Q: {values[q_col]}\nA: {values[a_col]}"
                extras = [f"{k}: {v}" for k, v in values.items() if k not in (q_col, a_col)]
                if extras:
                    chunk += "\n" + "\n".join(extras)
            else:
                chunk = "  |  ".join(f"{k}: {v}" for k, v in values.items())
            if chunk.strip():
                chunks.append(chunk.strip())
    return chunks


# ══════════════════════════════════════════════════════════════════════════════
# PDF PARSER — heading-aware fallback
# ══════════════════════════════════════════════════════════════════════════════

def _chunks_from_pdf(file_bytes: bytes) -> list[str]:
    import fitz
    doc  = fitz.open(stream=file_bytes, filetype="pdf")
    text = "\n".join(page.get_text() for page in doc)
    doc.close()

    lines    = text.split("\n")
    chunks   = []
    buf_head = ""
    buf_body: list[str] = []

    def is_heading(line: str) -> bool:
        s = line.strip()
        if not s:
            return False
        if re.match(r"^Q\d*[\s\-\u2013:]+", s, re.IGNORECASE):
            return True
        if s.isupper() and 3 < len(s) < 80:
            return True
        if s.endswith("?") and len(s) < 120:
            return True
        return False

    def flush():
        body = " ".join(buf_body).strip()
        if buf_head and body:
            chunks.append(f"{buf_head}\n{body}")
        elif buf_head:
            chunks.append(buf_head)
        elif body:
            chunks.append(body)

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if is_heading(stripped):
            flush(); buf_head = stripped; buf_body = []
        else:
            buf_body.append(stripped)

    flush()
    if not chunks or len(chunks) < 2:
        chunks = _word_window_chunks(text)
    return [c for c in chunks if len(c.strip()) > 15]


# ══════════════════════════════════════════════════════════════════════════════
# WORD-WINDOW FALLBACK
# ══════════════════════════════════════════════════════════════════════════════

def _word_window_chunks(text: str, chunk_size: int = 400, overlap: int = 50) -> list[str]:
    words  = text.split()
    chunks = []
    start  = 0
    while start < len(words):
        end   = min(start + chunk_size, len(words))
        chunk = " ".join(words[start:end])
        if len(chunk.strip()) > 20:
            chunks.append(chunk.strip())
        start += chunk_size - overlap
    return chunks


# ── Public dispatcher ──────────────────────────────────────────────────────────

def chunk_file(file_bytes: bytes, file_type: str) -> list[str]:
    ft = file_type.lower()
    if ft == "json":
        return _chunks_from_json(file_bytes)
    elif ft in ("md", "txt"):
        return _chunks_from_markdown(file_bytes.decode("utf-8", errors="ignore"))
    elif ft in ("docx", "doc"):
        return _chunks_from_docx(file_bytes)
    elif ft in ("xlsx", "xls"):
        return _chunks_from_xlsx(file_bytes)
    elif ft == "pdf":
        return _chunks_from_pdf(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def chunk_text(text: str, chunk_size: int = 400, overlap: int = 50) -> list[str]:
    """Backward-compat alias."""
    return _chunks_from_markdown(text)


# ── Embedding + FAISS ──────────────────────────────────────────────────────────

def embed_and_store(chunks: list[str], document_id: int) -> list[int]:
    import db
    model      = _get_model()
    idx        = _get_faiss_index()
    embeddings = model.encode(chunks, show_progress_bar=False, normalize_embeddings=True)
    embeddings = np.array(embeddings, dtype=np.float32)
    start_id   = idx.ntotal
    idx.add(embeddings)
    embedding_indices = list(range(start_id, start_id + len(chunks)))
    _embedding_store.extend(embeddings)
    _save_index()
    for i, (chunk, emb_idx) in enumerate(zip(chunks, embedding_indices)):
        db.insert_chunk(document_id, chunk, i, emb_idx)
    return embedding_indices


def ingest_file(file_bytes: bytes, filename: str) -> dict:
    import db
    ext    = filename.rsplit(".", 1)[-1].lower()
    title  = filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").title()
    chunks = chunk_file(file_bytes, ext)
    if not chunks:
        raise ValueError("No usable chunks could be extracted from this file.")
    doc_id            = db.insert_document(title, filename, ext)
    embedding_indices = embed_and_store(chunks, doc_id)
    return {"doc_id": doc_id, "title": title, "chunks": len(chunks), "embedding_indices": embedding_indices}


# ── Search helpers ─────────────────────────────────────────────────────────────

def keyword_search(query: str, chunks: list[dict], top_k: int = 5) -> list[dict]:
    query_terms = set(re.findall(r"\w+", query.lower()))
    scored = [(len(query_terms & set(re.findall(r"\w+", c["content"].lower()))), c)
              for c in chunks if query_terms & set(re.findall(r"\w+", c["content"].lower()))]
    scored.sort(key=lambda x: x[0], reverse=True)
    return [c for _, c in scored[:top_k]]


def semantic_search(query: str, top_k: int = 5) -> list[dict]:
    import db
    idx = _get_faiss_index()
    if idx.ntotal == 0:
        return []
    model       = _get_model()
    q_emb       = np.array(model.encode([query], normalize_embeddings=True), dtype=np.float32)
    distances, indices = idx.search(q_emb, min(top_k, idx.ntotal))
    emb_ids     = [int(i) for i in indices[0] if i >= 0]
    chunks      = db.get_chunks_by_ids(emb_ids)
    dist_map    = {int(indices[0][i]): float(distances[0][i]) for i in range(len(indices[0]))}
    for chunk in chunks:
        chunk["score"] = 1 / (1 + dist_map.get(chunk["embedding_index"], 999))
    chunks.sort(key=lambda x: x["score"], reverse=True)
    return chunks


def hybrid_search(query: str, top_k: int = 5) -> list[dict]:
    import db
    semantic = semantic_search(query, top_k=top_k)
    keyword  = keyword_search(query, db.get_all_chunks(), top_k=top_k)
    seen, merged = set(), []
    for chunk in semantic:
        if chunk["id"] not in seen:
            chunk["source"] = "semantic"; merged.append(chunk); seen.add(chunk["id"])
    for chunk in keyword:
        if chunk["id"] not in seen:
            chunk["source"] = "keyword"; chunk["score"] = chunk.get("score", 0.5)
            merged.append(chunk); seen.add(chunk["id"])
    return merged[:top_k]
